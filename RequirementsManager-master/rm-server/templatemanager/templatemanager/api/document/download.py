import os

from flask import request, send_file
from templatemanager.app import app
from templatemanager.dao.document import (
    Document, DocumentMongoDBDao
)
from templatemanager.mongodb import document_collection
from templatemanager.utils.handle_api import handle_response, handle_download

from urllib.parse import quote

import docx
import io

import base64

META_SUCCESS = {'status': 200, 'msg': '下载成功！'}
META_ERROR_NO_FILE = {'status': 404, 'msg': '下载失败，文档不存在！'}


# https://github.com/huyu1994/flask_create_word
@app.route('/document/download', methods=['GET'])
@handle_download
def document_download():
    # body = request.json

    document_id = request.args.get('document_id')

    document_mongodb_dao = DocumentMongoDBDao(document_collection)

    document = document_mongodb_dao.get_document(document_id)

    if not document:
        return {'meta': META_ERROR_NO_FILE}

    # generate the word.docx file
    word_docx = docx.Document()
    word_docx.add_heading(document.document_name, 1)

    # for outline, para in document.contents:
    #     word_docx.add_heading(outline, 3)
    #     word_docx.add_paragraph(para)
    if len(document.outline) != len(document.contents):
        print('wrong document file')
    i = 0
    while i < min(len(document.outline), len(document.contents)):
        word_docx.add_heading(document.outline[i], 3)
        word_docx.add_paragraph(document.contents[i])
        i += 1

    # save the file to a file stream
    docx_file = io.BytesIO()
    word_docx.save(docx_file)
    # rv = send_file(docx_file, mimetype='application/msword', as_attachment=True,
    #                attachment_filename=quote(document.document_name + '.docx'))
    # print(type(base64.b64encode(docx_file.getvalue()).decode('utf-8')))
    return {
        'meta': META_SUCCESS,
        'data': {
            'file_base64': base64.b64encode(docx_file.getvalue()).decode('utf-8')
        }
    }
    # response_info = {
    #     'meta': META_SUCCESS,
    #     'data': {
    #         'file_name': ,
    #         'file_buffer': docx_file.getvalue()
    #     }
    # }
    # return response_info


# if __name__ == '__main__':
#     word_docx = docx.Document()
#     word_docx.add_heading("test", 1)
#
#     # save the file to a file stream
#     docx_file = io.BytesIO()
#     word_docx.save(docx_file)
#     print(base64.b64encode(docx_file.getvalue()).decode('utf-8'))