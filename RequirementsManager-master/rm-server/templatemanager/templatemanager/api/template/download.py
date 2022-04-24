from flask import request
from templatemanager.app import app
from templatemanager.dao.template import (
    Template, TemplateMongoDBDao
)
from templatemanager.mongodb import template_collection
from templatemanager.utils.handle_api import handle_response

from docx import Document
import io

META_SUCCESS = {'status': 200, 'msg': '模板生成成功！'}
META_ERROR_NOT_EXIST = {'status': 404, 'msg': '生成失败，该模板不存在！'}


@app.route('/template/download', methods=[''])
@handle_response
def template_download():
    body = request.json

    template_name = body
    template_mongodb_dao = TemplateMongoDBDao(template_collection)
    # check if template not exists
    template = template_mongodb_dao.get_template(template_name)
    if not template:
        return {'meta': META_ERROR_NOT_EXIST}

    template_word_docx = Document()
    template_word_docx.add_heading(template.template_name, 0)
    for para in template.outline:
        template_word_docx.add_paragraph(para)

    # save the file to a file stream
    docx_file = io.FileIO()
    template_word_docx.save(docx_file)

    ret_info = {
        'meta': META_SUCCESS,
        'file': docx_file
    }
    docx_file.close()

    return {
        'meta': META_SUCCESS,
    }
